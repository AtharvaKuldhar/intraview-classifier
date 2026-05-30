"""
Technical Requirements Document Generator
=========================================
Generates a highly polished, professional, and beautifully formatted Microsoft Word (.docx)
technical requirements document from scratch, incorporating all verified clinical metrics,
model specifications, data preprocessing details, and cross-validation sweeps.

Uses best practices in typography and layout design:
- Harmonious color scheme (Executive Navy #1F4E79, Muted Gray #595959)
- Elegant tables with custom cell shading, borders, and column alignments
- Consistent hierarchical typography (Title, Subtitle, Headings, Styled Body)
- No double-spacing or carriage return hacks (uses paragraph formats)

Usage:
    python scripts/generate_formatted_requirements.py
"""
import os
import sys
import subprocess

# Auto-install python-docx if not present
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("[DOCX-GEN] python-docx not found. Installing now...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

# Colors definition (RGB)
COLOR_PRIMARY_RGB = RGBColor(31, 78, 121)     # Executive Navy (#1F4E79)
COLOR_SECONDARY_RGB = RGBColor(89, 89, 89)    # Slate Gray (#595959)
COLOR_TEXT_RGB = RGBColor(43, 43, 43)         # Charcoal (#2B2B2B)
COLOR_MUTED_RGB = RGBColor(128, 128, 128)     # Muted Gray

# XML Shading and Borders
XML_HEADER_SHADING = f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>'
XML_ALT_ROW_SHADING = f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>'
XML_CELL_PADDING = (
    f'<w:tcMar {nsdecls("w")}>'
    f'  <w:top w:w="120" w:type="dxa"/>'
    f'  <w:bottom w:w="120" w:type="dxa"/>'
    f'  <w:left w:w="150" w:type="dxa"/>'
    f'  <w:right w:w="150" w:type="dxa"/>'
    f'</w:tcMar>'
)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets internal padding (margins) inside table cells for highly readable layouts."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, xml_shading):
    """Sets background shading of a cell."""
    shading_elm = parse_xml(xml_shading)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    """Applies elegant, thin, gray horizontal borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1F4E79"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_run(run, font_name="Calibri", size_pt=11, color_rgb=COLOR_TEXT_RGB, bold=False, italic=False):
    """Applies clean typography to a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color_rgb
    run.font.bold = bold
    run.font.italic = italic

def add_heading_1(doc, text):
    """Adds a beautiful primary heading (Heading 1) in Executive Navy."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, font_name="Calibri Light", size_pt=16, color_rgb=COLOR_PRIMARY_RGB, bold=True)
    return p

def add_heading_2(doc, text):
    """Adds a beautiful secondary heading (Heading 2) in Slate Gray."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, font_name="Calibri", size_pt=13, color_rgb=COLOR_SECONDARY_RGB, bold=True)
    return p

def add_styled_paragraph(doc, text="", space_after=6, bold_prefix="", bullet=False):
    """Adds a consistently styled body paragraph with precise paragraph spacing."""
    style_name = 'List Bullet' if bullet else 'Normal'
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        format_run(r_prefix, font_name="Calibri", size_pt=11, color_rgb=COLOR_TEXT_RGB, bold=True)
        
    if text:
        r_text = p.add_run(text)
        format_run(r_text, font_name="Calibri", size_pt=11, color_rgb=COLOR_TEXT_RGB)
        
    return p

def main():
    print("[DOCX-GEN] Starting construction of clean requirements document...")
    doc = docx.Document()
    
    # 1. Page Margins Setup (Standard 1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # 2. Main Title (Centered)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Intraoral Image Classification Using Deep Learning")
    format_run(run_title, font_name="Calibri Light", size_pt=24, color_rgb=COLOR_PRIMARY_RGB, bold=True)
    
    # Subtitle (Centered)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Technical System Specifications & Rigorous Performance Requirements")
    format_run(run_sub, font_name="Calibri", size_pt=12, color_rgb=COLOR_SECONDARY_RGB, italic=True)
    
    # 3. Document Context / Preface
    add_styled_paragraph(
        doc,
        "This document details the complete system specifications, clinical data attributes, "
        "reproducible preprocessing methodologies, multi-architecture hyperparameters, hold-out comparative leaderboards, "
        "and multi-resolution K-Fold cross-validation robustness profiles. These requirements form the standardized "
        "framework for the intraoral dental view classification study, utilizing PyTorch and the Albumentations library."
    )
    
    # ==========================================================================
    # SECTION 1: CLINICAL DATASET CHARACTERISTICS
    # ==========================================================================
    add_heading_1(doc, "1. Clinical Dataset & Structural Attributes")
    add_styled_paragraph(
        doc,
        "Our analysis operates on a standardized, deduplicated, and class-aware offline augmented clinical dataset. "
        "Below are the verified structural parameters of the training, validation, and hold-out test divisions."
    )
    
    # Dataset Table
    dataset_table = doc.add_table(rows=6, cols=2)
    dataset_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(dataset_table)
    
    # Header
    set_cell_background(dataset_table.rows[0].cells[0], XML_HEADER_SHADING)
    set_cell_background(dataset_table.rows[0].cells[1], XML_HEADER_SHADING)
    dataset_table.rows[0].cells[0].text = "Attribute Field"
    dataset_table.rows[0].cells[1].text = "Technical Specification & Clinical Breakdown"
    
    # Data
    dataset_rows = [
        ("Total Number of Images", "5,234 balanced images (expanded from 2,975 unique cleaned instances)"),
        ("Intraoral Image Varieties", "8 Classifications: lower_front (650) | lower_left (684) | lower_occlusal (650) | lower_right (650) | upper_front (650) | upper_left (650) | upper_occlusal (650) | upper_right (650)"),
        ("Spatial Resolution", "Standardized at 224 x 224 pixels (RGB, 3 color channels)"),
        ("Dataset Split Ratio", "Stratified Partition: 70% Training / 15% Validation / 15% Hold-out Testing"),
        ("Cohort Split Cardinality", "3,663 training images, 785 validation images, and 786 unseen testing images")
    ]
    
    for idx, (attr, spec) in enumerate(dataset_rows, 1):
        row = dataset_table.rows[idx]
        row.cells[0].text = attr
        row.cells[1].text = spec
        if idx % 2 == 0:
            set_cell_background(row.cells[0], XML_ALT_ROW_SHADING)
            set_cell_background(row.cells[1], XML_ALT_ROW_SHADING)
            
    # Apply cell padding & typography to table cells
    for row_idx, row in enumerate(dataset_table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    if row_idx == 0:
                        format_run(run, font_name="Calibri", size_pt=11, color_rgb=RGBColor(255, 255, 255), bold=True)
                    else:
                        is_bold = (cell == row.cells[0])
                        format_run(run, font_name="Calibri", size_pt=10.5, color_rgb=COLOR_TEXT_RGB, bold=is_bold)
                        
    doc.add_paragraph().paragraph_format.space_after = Pt(6) # Spacing after table
    
    # ==========================================================================
    # SECTION 2: DATA PREPROCESSING PIPELINE
    # ==========================================================================
    add_heading_1(doc, "2. Data Preprocessing & Rigorous Integrity Pipeline")
    add_styled_paragraph(
        doc,
        "To enforce clinical hygiene, avoid overfitting, and eliminate label noise, a robust four-stage "
        "preprocessing pipeline was implemented:"
    )
    add_styled_paragraph(
        doc,
        bold_prefix="1. Cryptographic MD5 Deduplication: ",
        text="All raw images were scanned using MD5 checksum hashing. This eliminated 967 redundant duplicate "
             "within-class image copies (24.5% of raw collection). Additionally, 4 cross-class label conflicts "
             "(representing 8 image files) were completely discarded as label noise to ensure clear decision boundaries.",
        bullet=True
    )
    add_styled_paragraph(
        doc,
        bold_prefix="2. Dimensional Resolution Standardization: ",
        text="The remaining 2,975 unique cleaned images were structurally downsampled/upsampled using bilinear "
             "interpolation to a standard input resolution of 224 x 224 pixels with 3 RGB color channels.",
        bullet=True
    )
    add_styled_paragraph(
        doc,
        bold_prefix="3. Tier-Based Class-Aware Augmentation: ",
        text="Offline augmentations via Albumentations balanced all minority classes to a targeted threshold of "
             "~650 images each, producing 2,259 unique synthetic images. Modalities were limited to small rotations "
             "(max +-15 degrees), exposure/contrast jitter, random resized crop, Gaussian defocus blur, tissue saturation "
             "value adjustment (Hue locked at 0), and sharpening. In-memory MD5 unique filters guaranteed zero exact duplicate "
             "outputs in the final balanced dataset of 5,234 images. All flips (horizontal/vertical) and elastic warping "
             "were strictly banned to prevent Left-Right view confusion and clinical anatomical distortion.",
        bullet=True
    )
    add_styled_paragraph(
        doc,
        bold_prefix="4. Isolated Stratified Splitting: ",
        text="The final balanced cohort of 5,234 images was partitioned into 70% Training (3,663), 15% Validation (785), "
             "and 15% Hold-out Testing (786) to ensure unbiased generalizability metrics.",
        bullet=True
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================================================
    # SECTION 3: SYSTEM HYPERPARAMETERS COMPARISON
    # ==========================================================================
    add_heading_1(doc, "3. Hyperparameters & Architectural Configurations")
    add_styled_paragraph(
        doc,
        "Five distinct paradigm families (Classical CNN, Dense CNN, Lightweight CNN, Modern CNN, and Vision Transformer) "
        "were evaluated under identical conditions. Below is the detailed comparative hyperparameter configuration matrix."
    )
    
    hparams_table = doc.add_table(rows=18, cols=6)
    hparams_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(hparams_table)
    
    # Headers
    hparams_headers = ["Hyperparameter", "ResNet-50", "DenseNet-121", "MobileNetV3", "ConvNeXt-Tiny", "Swin-Tiny"]
    for idx, h in enumerate(hparams_headers):
        cell = hparams_table.rows[0].cells[idx]
        set_cell_background(cell, XML_HEADER_SHADING)
        cell.text = h
        
    # Specs data
    hparams_rows = [
        ("Model Backbone", "resnet50", "densenet121", "mobilenetv3_small", "convnext_tiny", "swin_tiny"),
        ("Model Initialization", "Pretrained ImageNet-1K", "Pretrained ImageNet-1K", "Pretrained ImageNet-1K", "Pretrained ImageNet-1K", "Pretrained ImageNet-22K"),
        ("Convolution Layers", "50 layers deep, residual", "121 layers deep, dense", "Depthwise separable", "Modernized blocks", "None (Vision Transformer)"),
        ("Attention Mechanism", "None", "None", "Squeeze-and-Excitation", "None", "Shifted Window Self-Attn"),
        ("Pooling Strategy", "Global Average Pooling", "Global Average Pooling", "Global Average Pooling", "Global Average Pooling", "Global Average Pooling"),
        ("FC Output Layer", "Linear (768 -> 8)", "Linear (1024 -> 8)", "Linear (1024 -> 8)", "Linear (768 -> 8)", "Linear (768 -> 8)"),
        ("Activation Function", "ReLU", "ReLU", "Hard-Swish", "GELU", "GELU"),
        ("Output Activation", "Softmax", "Softmax", "Softmax", "Softmax", "Softmax"),
        ("Loss Function", "Weighted Cross-Entropy", "Weighted Cross-Entropy", "Weighted Cross-Entropy", "Weighted Cross-Entropy", "Weighted Cross-Entropy"),
        ("Optimizer Type", "AdamW (decay=0.01)", "AdamW (decay=0.01)", "AdamW (decay=0.01)", "AdamW (decay=0.01)", "AdamW (decay=0.01)"),
        ("Learning Rate Strategy", "P1: 1e-3, P2: 1e-4", "P1: 1e-3, P2: 1e-4", "P1: 1e-3, P2: 1e-4", "P1: 1e-3, P2: 1e-4", "P1: 1e-3, P2: 1e-4"),
        ("Batch Size", "32", "32", "32", "32", "32"),
        ("Number of Epochs", "50 max (10 P1, 40 P2)", "50 max (10 P1, 40 P2)", "50 max (10 P1, 40 P2)", "50 max (10 P1, 40 P2)", "50 max (10 P1, 40 P2)"),
        ("Dropout / DropPath", "0.0 (None)", "0.0 (None)", "0.2 (Dropout)", "0.0 (None)", "0.2 (Stochastic Depth)"),
        ("Data Augmentation", "Tiered Offline Albumentations", "Tiered Offline Albumentations", "Tiered Offline Albumentations", "Tiered Offline Albumentations", "Tiered Offline Albumentations"),
        ("Fine-Tuning gradual", "Two-stage unfreezing", "Two-stage unfreezing", "Two-stage unfreezing", "Two-stage unfreezing", "Two-stage unfreezing"),
        ("Evaluation Metrics", "Acc, Prec, Recall, F1", "Acc, Prec, Recall, F1", "Acc, Prec, Recall, F1", "Acc, Prec, Recall, F1", "Acc, Prec, Recall, F1")
    ]
    
    for row_idx, data in enumerate(hparams_rows, 1):
        row = hparams_table.rows[row_idx]
        for col_idx, text in enumerate(data):
            row.cells[col_idx].text = text
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, XML_ALT_ROW_SHADING)
                
    # Style cells
    for row_idx, row in enumerate(hparams_table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=100, bottom=100) # Slightly tighter margins due to large row count
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    if row_idx == 0:
                        format_run(run, font_name="Calibri", size_pt=10.5, color_rgb=RGBColor(255, 255, 255), bold=True)
                    else:
                        is_bold = (col_idx == 0)
                        format_run(run, font_name="Calibri", size_pt=9.5, color_rgb=COLOR_TEXT_RGB, bold=is_bold)
                        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================================================
    # SECTION 4: CROSS-ARCHITECTURE PERFORMANCE LEADERBOARD
    # ==========================================================================
    add_heading_1(doc, "4. Cross-Architecture Comparative Evaluation Leaderboard")
    add_styled_paragraph(
        doc,
        "The trained models were evaluated on the strictly isolated 15% stratified hold-out test set "
        "(786 unseen images). The leaderboard below is ranked descending by F1-Score (Macro), representing "
        "the standard metric for class-imbalance robustness in diagnostic medical imaging classification."
    )
    
    perf_table = doc.add_table(rows=9, cols=8)
    perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(perf_table)
    
    perf_headers = ["Model Architecture", "Paradigm Family", "Params", "FLOPs", "Accuracy", "Precision", "Recall", "F1-Score"]
    for idx, h in enumerate(perf_headers):
        cell = perf_table.rows[0].cells[idx]
        set_cell_background(cell, XML_HEADER_SHADING)
        cell.text = h
        
    perf_rows = [
        ("Swin-Tiny", "Vision Transformer", "28.3M", "4.5G", "99.11%", "99.13%", "99.12%", "99.11%"),
        ("ConvNeXt-Tiny", "Modern CNN", "28.6M", "4.5G", "99.11%", "99.11%", "99.11%", "99.11%"),
        ("MobileNetV3-Small", "Efficient CNN", "2.5M", "0.06G", "98.98%", "98.99%", "99.00%", "98.99%"),
        ("DenseNet-121", "Classical CNN", "8.0M", "2.9G", "98.98%", "98.98%", "98.98%", "98.98%"),
        ("ResNet-50", "Classical CNN", "25.6M", "4.1G", "98.85%", "98.86%", "98.85%", "98.85%"),
        ("EfficientNet-B2", "Efficient CNN", "9.1M", "1.0G", "98.35%", "98.37%", "98.35%", "98.35%"),
        ("EfficientNet-B3", "Efficient CNN", "12.2M", "1.8G", "97.46%", "97.48%", "97.46%", "97.46%"),
        ("DINOv2-Small (ViT-S/14)", "Vision Transformer", "22.1M", "4.6G", "94.78%", "95.21%", "94.77%", "94.81%")
    ]
    
    for row_idx, data in enumerate(perf_rows, 1):
        row = perf_table.rows[row_idx]
        for col_idx, text in enumerate(data):
            row.cells[col_idx].text = text
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, XML_ALT_ROW_SHADING)
                
    # Style cells
    for row_idx, row in enumerate(perf_table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=110, bottom=110)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    if row_idx == 0:
                        format_run(run, font_name="Calibri", size_pt=10.5, color_rgb=RGBColor(255, 255, 255), bold=True)
                    else:
                        is_bold = (col_idx == 0 or row_idx == 1) # Bold model names & the champion row
                        color = COLOR_PRIMARY_RGB if row_idx == 1 else COLOR_TEXT_RGB
                        format_run(run, font_name="Calibri", size_pt=9.5, color_rgb=color, bold=is_bold)
                        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================================================
    # SECTION 5: K-FOLD CROSS-VALIDATION ROBUSTNESS SWEEP
    # ==========================================================================
    add_heading_1(doc, "5. Statistical Robustness & K-Fold Cross-Validation Sweeps")
    add_styled_paragraph(
        doc,
        "To validate that the peak generalizability of our champion model (Swin-Tiny) was not an artifact of "
        "a specific train-test partition, we subjected it to a rigorous multi-resolution stratified cross-validation "
        "sweep across K in {2, 3, 5, 7}. The consistently tight standard deviations across all configurations "
        "prove that the model is highly invariant to data splitting and represents an clinically-stable classifier."
    )
    
    kfold_table = doc.add_table(rows=6, cols=6)
    kfold_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(kfold_table)
    
    kfold_headers = ["Dataset Identifier", "K-Fold Configuration", "Accuracy % (μ ± σ)", "Precision % (μ ± σ)", "Recall % (μ ± σ)", "F1-Score % (μ ± σ)"]
    for idx, h in enumerate(kfold_headers):
        cell = kfold_table.rows[0].cells[idx]
        set_cell_background(cell, XML_HEADER_SHADING)
        cell.text = h
        
    kfold_rows = [
        ("Intraoral dataset", "K=2 (50% train split)", "98.18% ± 0.10%", "98.20% ± 0.08%", "98.19% ± 0.10%", "98.18% ± 0.10%"),
        ("Intraoral dataset", "K=3 (66% train split)", "99.04% ± 0.22%", "99.05% ± 0.22%", "99.05% ± 0.22%", "99.04% ± 0.22%"),
        ("Intraoral dataset", "K=5 (80% train split)", "99.12% ± 0.29%", "99.14% ± 0.28%", "99.12% ± 0.29%", "99.12% ± 0.29%"),
        ("Intraoral dataset", "K=7 (85.7% train split)", "99.39% ± 0.17%", "99.40% ± 0.17%", "99.39% ± 0.17%", "99.39% ± 0.17%"),
        ("Intraoral dataset", "K=9 (88.9% train split)", "N/A (Not Evaluated)", "N/A", "N/A", "N/A")
    ]
    
    for row_idx, data in enumerate(kfold_rows, 1):
        row = kfold_table.rows[row_idx]
        for col_idx, text in enumerate(data):
            row.cells[col_idx].text = text
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, XML_ALT_ROW_SHADING)
                
    # Style cells
    for row_idx, row in enumerate(kfold_table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=110, bottom=110)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    if row_idx == 0:
                        format_run(run, font_name="Calibri", size_pt=10.5, color_rgb=RGBColor(255, 255, 255), bold=True)
                    else:
                        is_bold = (col_idx == 1)
                        format_run(run, font_name="Calibri", size_pt=9.5, color_rgb=COLOR_TEXT_RGB, bold=is_bold)
                        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================================================
    # SECTION 6: SYSTEM FLOW DIAGRAM
    # ==========================================================================
    add_heading_1(doc, "6. Technical System Execution Flow")
    add_styled_paragraph(
        doc,
        "The comprehensive system pipeline represents a unified engineering workflow, running from raw clinical "
        "ingestion to publication comparison output:"
    )
    
    # Custom quote box for system flow
    flow_table = doc.add_table(rows=1, cols=1)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(flow_table.rows[0].cells[0], top=150, bottom=150, left=200, right=200)
    set_cell_background(flow_table.rows[0].cells[0], XML_ALT_ROW_SHADING)
    
    # Left border color highlight
    tcPr = flow_table.rows[0].cells[0]._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="1F4E79"/>'
        f'  <w:right w:val="none"/>'
    '</w:tcBorders>'
    )
    tcPr.append(borders)
    
    system_flow_text = (
        "Clinical Image Dataset Ingestion ➔ Cryptographic MD5 Deduplication (Phase 1) ➔ "
        "Bilinear Standardized Resizing (224x224 RGB) ➔ Class-Aware Tiered Offline Augmentation (Phase 2) ➔ "
        "Stratified Train/Val/Test Isolation (70/15/15) ➔ Double-Phase Optimizer Training (Phase 1 Head Calibration ➔ "
        "Phase 2 Global Fine-Tuning) ➔ Hold-out Leaderboard Benchmarking ➔ Multi-Resolution K-Fold Sweeps "
        "(Statistical Proofs) ➔ Publication-Grade Comparison Report Compilation"
    )
    
    flow_cell_p = flow_table.rows[0].cells[0].paragraphs[0]
    flow_cell_p.paragraph_format.line_spacing = 1.25
    flow_cell_p.paragraph_format.space_after = Pt(0)
    run_flow = flow_cell_p.add_run(system_flow_text)
    format_run(run_flow, font_name="Calibri", size_pt=10, color_rgb=COLOR_PRIMARY_RGB, bold=True)
    
    # Save document
    output_dir = "data"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Requirements_populated_formatted.docx")
    doc.save(output_path)
    print(f"[DOCX-GEN] SUCCESS: Publication-grade document generated at: '{output_path}'")
    
if __name__ == "__main__":
    main()
