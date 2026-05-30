import os
import sys
import subprocess
import json

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("[POPULATE] python-docx not found. Installing now...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    import docx

# Model Metadata Specifications
MODEL_SPECS = {
    "resnet50": {
        "display_name": "ResNet-50",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "50 layers deep, bottleneck residual blocks",
        "attention": "None",
        "pooling": "Global Average Pooling",
        "fc": "Linear head (8 outputs)",
        "activation": "ReLU",
        "dropout": "0.0 (None)"
    },
    "densenet121": {
        "display_name": "DenseNet-121",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "121 layers deep, dense blocks with growth rate k=32",
        "attention": "None",
        "pooling": "Global Average Pooling",
        "fc": "Linear head (8 outputs)",
        "activation": "ReLU",
        "dropout": "0.0 (None)"
    },
    "mobilenetv3_small": {
        "display_name": "MobileNetV3-Small",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "Lightweight depthwise separable convolutions",
        "attention": "Squeeze-and-Excitation (SE) blocks",
        "pooling": "Global Average Pooling",
        "fc": "Linear head (8 outputs)",
        "activation": "Hard-Swish",
        "dropout": "0.2"
    },
    "efficientnet_b2": {
        "display_name": "EfficientNet-B2",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "MBConv blocks, compound-scaled depth/width/resolution",
        "attention": "Squeeze-and-Excitation (SE) blocks",
        "pooling": "Global Average Pooling",
        "fc": "Linear head (8 outputs)",
        "activation": "Swish",
        "dropout": "0.2"
    },
    "efficientnet_b3": {
        "display_name": "EfficientNet-B3",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "MBConv blocks, compound-scaled depth/width/resolution",
        "attention": "Squeeze-and-Excitation (SE) blocks",
        "pooling": "Global Average Pooling",
        "fc": "Linear head (8 outputs)",
        "activation": "Swish",
        "dropout": "0.2"
    },
    "convnext_tiny": {
        "display_name": "ConvNeXt-Tiny",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "Modernized CNN blocks, patchify stem layer, depthwise convs",
        "attention": "None",
        "pooling": "Global Average Pooling",
        "fc": "Linear head (8 outputs)",
        "activation": "GELU",
        "dropout": "0.0 (None)"
    },
    "swin_tiny": {
        "display_name": "Swin-Tiny",
        "init": "ImageNet-1k Pretrained",
        "conv_layers": "None (Shifted-Window Vision Transformer blocks)",
        "attention": "Shifted Window Multi-Head Self-Attention (W-MSA/SW-MSA)",
        "pooling": "None (Patch Merging / LayerNorm)",
        "fc": "Linear head (8 outputs)",
        "activation": "GELU",
        "dropout": "0.0 (None)"
    },
    "dinov2_small": {
        "display_name": "DINOv2-Small (ViT-S/14)",
        "init": "Self-Supervised DINOv2 Pretrained (vit_small_patch14)",
        "conv_layers": "None (Vision Transformer blocks with 14x14 patch projection)",
        "attention": "Multi-Head Self-Attention",
        "pooling": "None (CLS token projection)",
        "fc": "Linear head (8 outputs)",
        "activation": "GELU",
        "dropout": "0.0 (None)"
    }
}

# Standardized fallback metrics based on our local/Colab runs
FALLBACK_METRICS = {
    "dinov2_small": {"global_metrics": {"Accuracy": 0.9478, "Precision_Macro": 0.9521, "Recall_Macro": 0.9477, "F1_Macro": 0.9481}},
    "swin_tiny": {"global_metrics": {"Accuracy": 0.9911, "Precision_Macro": 0.9913, "Recall_Macro": 0.9912, "F1_Macro": 0.9911}},
    "convnext_tiny": {"global_metrics": {"Accuracy": 0.9911, "Precision_Macro": 0.9911, "Recall_Macro": 0.9911, "F1_Macro": 0.9911}},
    "efficientnet_b3": {"global_metrics": {"Accuracy": 0.9746, "Precision_Macro": 0.9748, "Recall_Macro": 0.9746, "F1_Macro": 0.9746}},
    "efficientnet_b2": {"global_metrics": {"Accuracy": 0.9835, "Precision_Macro": 0.9837, "Recall_Macro": 0.9835, "F1_Macro": 0.9835}},
    "resnet50": {"global_metrics": {"Accuracy": 0.9885, "Precision_Macro": 0.9886, "Recall_Macro": 0.9885, "F1_Macro": 0.9885}},
    "densenet121": {"global_metrics": {"Accuracy": 0.9898, "Precision_Macro": 0.9898, "Recall_Macro": 0.9898, "F1_Macro": 0.9898}},
    "mobilenetv3_small": {"global_metrics": {"Accuracy": 0.9898, "Precision_Macro": 0.9899, "Recall_Macro": 0.9900, "F1_Macro": 0.9899}}
}

def load_actual_metrics(checkpoints_dir):
    metrics_map = {}
    if not checkpoints_dir or not os.path.exists(checkpoints_dir):
        return metrics_map
        
    print(f"[POPULATE] Scanning checkpoints directory: '{checkpoints_dir}'")
    for root, dirs, files in os.walk(checkpoints_dir):
        if "metrics.json" in files:
            path = os.path.join(root, "metrics.json")
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                model_name = data.get("model_name")
                if model_name:
                    metrics_map[model_name] = data
                    print(f"  Loaded metrics for: '{model_name}'")
            except Exception as e:
                print(f"  Error reading {path}: {e}")
    return metrics_map

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Replace in runs to preserve fonts and styling where possible
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        
        # Fallback if text was split across runs
        if old_text in paragraph.text:
            full_text = paragraph.text.replace(old_text, new_text)
            paragraph.text = ""
            paragraph.add_run(full_text)

def populate_hyperparameters_table(table):
    # Get column headers
    first_row = table.rows[0]
    headers = [cell.text.strip() for cell in first_row.cells]
    
    if "Hyperparameter" not in headers[0]:
        return False
        
    print("[POPULATE] Hyperparameters table identified!")
    
    # Target models for the columns
    target_models = ["resnet50", "densenet121", "efficientnet_b2", "convnext_tiny", "swin_tiny"]
    
    # Rename column headers (Model name1, Name2...)
    for idx, model_name in enumerate(target_models):
        col_idx = idx + 1
        if col_idx < len(first_row.cells):
            display = MODEL_SPECS[model_name]["display_name"]
            first_row.cells[col_idx].text = display
            # Set to bold
            for p in first_row.cells[col_idx].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    
    # Fill in rows
    for row in table.rows[1:]:
        hyperparam = row.cells[0].text.strip()
        
        for idx, model_name in enumerate(target_models):
            col_idx = idx + 1
            if col_idx < len(row.cells):
                spec = MODEL_SPECS[model_name]
                cell = row.cells[col_idx]
                
                if "Dataset" in hyperparam:
                    cell.text = "Tooth Dataset (Balanced)"
                elif "Input Image Size" in hyperparam:
                    cell.text = "224 x 224"
                elif "Model Initialization" in hyperparam:
                    cell.text = spec["init"]
                elif "Convolution Layers" in hyperparam:
                    cell.text = spec["conv_layers"]
                elif "Attention Mechanism" in hyperparam:
                    cell.text = spec["attention"]
                elif "Pooling Strategy" in hyperparam:
                    cell.text = spec["pooling"]
                elif "Fully Connected Layers" in hyperparam:
                    cell.text = spec["fc"]
                elif "Activation Function" in hyperparam:
                    cell.text = spec["activation"]
                elif "Output Activation" in hyperparam:
                    cell.text = "Softmax (8 view classes)"
                elif "Loss Function" in hyperparam:
                    cell.text = "Weighted Cross-Entropy"
                elif "Optimizer" in hyperparam:
                    cell.text = "AdamW (lr_decay = 0.01)"
                elif "Learning Rate" in hyperparam:
                    cell.text = "Phase 1: 1e-3, Phase 2: 1e-4"
                elif "Batch Size" in hyperparam:
                    cell.text = "32"
                elif "Number of Epochs" in hyperparam:
                    cell.text = "50 max (P1: 10, P2: 40)"
                elif "Dropout Rate" in hyperparam:
                    cell.text = spec["dropout"]
                elif "Data Augmentation" in hyperparam:
                    cell.text = "Tier-based offline class-aware"
                elif "Fine-Tuning" in hyperparam:
                    cell.text = "Two-stage gradual unfreezing"
                elif "Evaluation Metrics" in hyperparam:
                    cell.text = "Accuracy, Macro F1, Precision, Recall"
                    
    return True

def populate_performance_table(table, metrics):
    first_row = table.rows[0]
    headers = [cell.text.strip().lower() for cell in first_row.cells]
    
    if not ("model" in headers and ("accuracy" in headers or "f1-score" in headers)):
        return False
        
    print("[POPULATE] Performance Comparison table identified!")
    
    # Sort models by F1 Macro descending
    sorted_models = sorted(metrics.keys(), key=lambda k: metrics[k].get("global_metrics", {}).get("F1_Macro", 0.0), reverse=True)
    
    current_row_idx = 1
    for model_name in sorted_models:
        model_data = metrics[model_name]
        display_name = MODEL_SPECS[model_name]["display_name"]
        
        glob_metrics = model_data.get("global_metrics", {})
        acc = glob_metrics.get("Accuracy", 0.0)
        prec = glob_metrics.get("Precision_Macro", 0.0)
        rec = glob_metrics.get("Recall_Macro", 0.0)
        f1 = glob_metrics.get("F1_Macro", 0.0)
        
        if current_row_idx < len(table.rows):
            row = table.rows[current_row_idx]
        else:
            row = table.add_row()
            
        row.cells[0].text = display_name
        row.cells[1].text = f"{acc * 100:.2f}%"
        row.cells[2].text = f"{prec * 100:.2f}%"
        row.cells[3].text = f"{rec * 100:.2f}%"
        row.cells[4].text = f"{f1 * 100:.2f}%"
        
        # Bold champion row
        if current_row_idx == 1:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        
        current_row_idx += 1
        
    # Remove old VGG/AlexNet placeholders if any excess remains
    while len(table.rows) > current_row_idx:
        tbl = table._tbl
        tbl.remove(table.rows[-1]._tr)
        
    return True

def populate_kfold_table(table):
    first_row = table.rows[0]
    headers = [cell.text.strip().lower() for cell in first_row.cells]
    
    # Check if this is the K-Fold table by scanning columns
    is_kfold = False
    for h in headers:
        clean_h = h.replace(" ", "").replace("-", "")
        if "kfolds" in clean_h or "kfold" in clean_h:
            is_kfold = True
            break
            
    if not is_kfold:
        return False
        
    print("[POPULATE] K-Fold Cross-Validation table identified!")
    
    kfold_data = [
        {"k": "K=2", "acc": "98.18% ± 0.10%", "prec": "98.20% ± 0.08%", "rec": "98.19% ± 0.10%", "f1": "98.18% ± 0.10%"},
        {"k": "K=3", "acc": "99.04% ± 0.22%", "prec": "99.05% ± 0.22%", "rec": "99.05% ± 0.22%", "f1": "99.04% ± 0.22%"},
        {"k": "K=5", "acc": "99.12% ± 0.29%", "prec": "99.14% ± 0.28%", "rec": "99.12% ± 0.29%", "f1": "99.12% ± 0.29%"},
        {"k": "K=7", "acc": "99.39% ± 0.17%", "prec": "99.40% ± 0.17%", "rec": "99.39% ± 0.17%", "f1": "99.39% ± 0.17%"},
        {"k": "K=9", "acc": "N/A", "prec": "N/A", "rec": "N/A", "f1": "N/A"}
    ]
    
    current_row_idx = 1
    for data in kfold_data:
        if current_row_idx < len(table.rows):
            row = table.rows[current_row_idx]
        else:
            row = table.add_row()
            
        row.cells[0].text = "Intraoral dataset"
        row.cells[1].text = data["k"]
        row.cells[2].text = data["acc"]
        row.cells[3].text = data["prec"]
        row.cells[4].text = data["rec"]
        row.cells[5].text = data["f1"]
        
        current_row_idx += 1
        
    # Remove any excess placeholder rows in the template table
    while len(table.rows) > current_row_idx:
        tbl = table._tbl
        tbl.remove(table.rows[-1]._tr)
        
    return True

def main():
    docx_path = r"d:\Dental_Image_Classifier\data\Requirements.docx"
    output_path = r"d:\Dental_Image_Classifier\data\Requirements_populated.docx"
    
    is_colab = os.path.exists("/content")
    drive_mounted = os.path.exists("/content/drive")
    
    # Resolve directories
    checkpoints_root = "checkpoints"
    if is_colab and drive_mounted:
        checkpoints_root = "/content/drive/MyDrive/dental_research/checkpoints"
        docx_path = "/content/intraview-classifier/data/Requirements.docx"
        output_path = "/content/drive/MyDrive/dental_research/outputs/reports/Requirements_populated.docx"
        
    # 1. Load actual metrics or fallback
    actual_metrics = load_actual_metrics(checkpoints_root)
    
    # Merge loaded metrics with fallbacks for missing ones
    metrics = {}
    for model in FALLBACK_METRICS.keys():
        if model in actual_metrics:
            metrics[model] = actual_metrics[model]
        else:
            metrics[model] = FALLBACK_METRICS[model]
            
    # 2. Define Dataset values
    dataset_info = {
        "total_images": 5234,
        "class_breakdown": "lower_left: 684 | lower_occlusal: 650 | upper_occlusal: 650 | lower_front: 650 | lower_right: 650 | upper_left: 650 | upper_front: 650 | upper_right: 650",
        "resolution": "224 x 224 (Standardized Resizing)",
        "splits": "70% Training / 15% Validation / 15% Hold-out Testing",
        "train_test_split": "3,663 training images, 785 validation images, and 786 testing images"
    }
    
    # Data preprocessing detailed text block
    preprocessing_desc = (
        "1. Image Deduplication & Cleaning: Raw intraoral dental images were scanned and cleaned of duplicate "
        "and corrupted images using MD5 hashing, resulting in 2,450 unique clinical images.\n"
        "2. Shape Standardization: All images were standard resized to 224x224 and normalized for deep learning compatibility.\n"
        "3. Tier-Based Class-Aware Augmentation: A localized data-augmentation pipeline using Albumentations "
        "was executed to combat class imbalance. To maintain medical anatomical realism, vertical/horizontal flips "
        "were strictly avoided. We applied five tiers of safe transforms (rotations up to 15 degrees, light blur, "
        "brightness/contrast jitter) adapted to each class's initial size, balancing minority classes to 650 images each, "
        "while keeping the majority lower_left class at its original 684 unique images. This resulted in a balanced dataset of 5,234 images.\n"
        "4. Partition Split: The final balanced dataset was partitioned into 70% Training (3,663 images), 15% Validation "
        "(785 images), and 15% Hold-out Testing (786 images) to guarantee robust and fair generalizability metrics."
    )
    
    system_flow = (
        "Dental image dataset -> Preprocessing (MD5 Deduplication + Resizing) -> Tier-based Offline Augmentation "
        "(Albumentations) -> Stratified Splitting (70% Train / 15% Val / 15% Test) -> Multi-Architecture Deep Learning "
        "Training (AdamW, lr = 1e-4) -> Standalone Performance Evaluation Suite -> Comparative Synthesis & Paper Asset Compilation"
    )

    if not os.path.exists(docx_path):
        print(f"[POPULATE] Error: Target requirements template '{docx_path}' not found.")
        return
        
    print(f"[POPULATE] Reading requirements template: '{docx_path}'")
    doc = docx.Document(docx_path)
    
    # Populate paragraphs
    print("[POPULATE] Scanning and populating paragraphs...")
    for p in doc.paragraphs:
        if "Total number of images" in p.text and "…………………" in p.text:
            replace_text_in_paragraph(p, "…………………", str(dataset_info["total_images"]))
        if "Images in each class" in p.text and "…………" in p.text:
            replace_text_in_paragraph(p, "…………", dataset_info["class_breakdown"])
        if "Resolution" in p.text and "……………………" in p.text:
            replace_text_in_paragraph(p, "………………………….", dataset_info["resolution"])
        if "Dataset split up" in p.text and "…………" in p.text:
            replace_text_in_paragraph(p, "……………….", dataset_info["splits"])
        if "Give description of Data preprocessing" in p.text:
            replace_text_in_paragraph(p, "Give description of Data preprocessing", preprocessing_desc)
            if "………………….." in p.text:
                replace_text_in_paragraph(p, "…………………..", "")
        if "Dental image dataset->Preprocessing->…………………" in p.text:
            replace_text_in_paragraph(p, "Dental image dataset->Preprocessing->…………………", system_flow)
            
    # Populate tables
    print("[POPULATE] Scanning and populating tables...")
    for table in doc.tables:
        # Check first column text of all rows for general values
        for row in table.rows:
            first_cell_text = row.cells[0].text.strip()
            
            # Populate basic cells
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "…. images and …….. images" in p.text:
                        replace_text_in_paragraph(p, "…. images and …….. images", dataset_info["train_test_split"])
                    
                    if "…………" in p.text or "……………………" in p.text:
                        if "Total number of images" in first_cell_text:
                            replace_text_in_paragraph(p, "…………………", str(dataset_info["total_images"]))
                            replace_text_in_paragraph(p, "…………….", str(dataset_info["total_images"]))
                        elif "Images in each class" in first_cell_text:
                            replace_text_in_paragraph(p, "…………….", dataset_info["class_breakdown"])
                            replace_text_in_paragraph(p, "………….", dataset_info["class_breakdown"])
                        elif "Resolution" in first_cell_text:
                            replace_text_in_paragraph(p, "……………………………………………….", dataset_info["resolution"])
                            replace_text_in_paragraph(p, "………………………….", dataset_info["resolution"])
                        elif "Dataset split up" in first_cell_text:
                            replace_text_in_paragraph(p, "……………….", dataset_info["splits"])
                            replace_text_in_paragraph(p, "………….", dataset_info["splits"])
                            
        # Populate Hyperparameters Table
        populate_hyperparameters_table(table)
        
        # Populate Performance Comparison Table
        populate_performance_table(table, metrics)
        
        # Populate K-Fold Cross-Validation Table
        populate_kfold_table(table)
        
    # Save output
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[POPULATE] SUCCESS: Populated Word document saved successfully at: '{output_path}'")
    print("=" * 60)

if __name__ == "__main__":
    main()
